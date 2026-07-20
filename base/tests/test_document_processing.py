import io
from unittest.mock import MagicMock, patch

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import SimpleTestCase

from base.document_processing import MAX_CHUNK_LENGTH, chunk_text, extract_text
from base.models import Document


class ExtractTextTests(SimpleTestCase):
    def test_txt(self):
        doc = Document(file=SimpleUploadedFile("notes.txt", b"Hello from a text file."))
        self.assertEqual(extract_text(doc), "Hello from a text file.")

    def test_pdf(self):
        doc = Document(file=SimpleUploadedFile("notes.pdf", b"%PDF-1.4 fake bytes"))
        fake_page_1 = MagicMock()
        fake_page_1.extract_text.return_value = "Page one text."
        fake_page_2 = MagicMock()
        fake_page_2.extract_text.return_value = "Page two text."
        with patch("base.document_processing.PdfReader") as mock_reader_cls:
            mock_reader_cls.return_value.pages = [fake_page_1, fake_page_2]
            result = extract_text(doc)
        self.assertEqual(result, "Page one text.\n\nPage two text.")

    def test_docx(self):
        from docx import Document as DocxDocument

        buffer = io.BytesIO()
        docx_doc = DocxDocument()
        docx_doc.add_paragraph("First paragraph.")
        docx_doc.add_paragraph("Second paragraph.")
        docx_doc.save(buffer)

        doc = Document(file=SimpleUploadedFile("notes.docx", buffer.getvalue()))
        result = extract_text(doc)
        self.assertEqual(result, "First paragraph.\n\nSecond paragraph.")

    def test_unsupported_extension_raises(self):
        doc = Document(file=SimpleUploadedFile("notes.xyz", b"whatever"))
        with self.assertRaises(ValueError):
            extract_text(doc)


class ChunkTextTests(SimpleTestCase):
    def test_splits_on_paragraph_breaks(self):
        text = (
            "First paragraph here with enough words to pass.\n\n"
            "Second paragraph here too, also long enough now."
        )
        self.assertEqual(
            chunk_text(text),
            [
                "First paragraph here with enough words to pass.",
                "Second paragraph here too, also long enough now.",
            ],
        )

    def test_drops_short_chunks(self):
        text = (
            "This is a long enough paragraph to keep.\n\n"
            "short\n\n"
            "Another long enough paragraph to keep."
        )
        result = chunk_text(text)
        self.assertEqual(len(result), 2)
        self.assertNotIn("short", result)

    def test_no_paragraph_breaks_returns_single_chunk(self):
        text = "Just one long paragraph with no breaks in it at all here."
        self.assertEqual(chunk_text(text), [text])

    def test_long_paragraph_splits_into_sentence_chunks(self):
        text = (
            "The company reported strong results across all business units this year. "
            "Revenue grew steadily due to increased customer demand in urban markets. "
            "Operating costs were kept under control through efficiency programs. "
            "Management expects continued growth in the coming fiscal year."
        )
        self.assertGreater(len(text), MAX_CHUNK_LENGTH)  # sanity check: fixture must trigger the split path

        result = chunk_text(text)

        self.assertGreater(len(result), 1)
        for chunk in result:
            self.assertLessEqual(len(chunk), MAX_CHUNK_LENGTH)
        joined = " ".join(result)
        self.assertIn("Revenue grew steadily", joined)
        self.assertIn("Management expects continued growth", joined)

    def test_single_long_sentence_kept_whole(self):
        text = (
            "This is one single very long run on sentence without any period "
            "in the middle that just keeps going and going through many many "
            "words until it finally reaches its very end without ever giving "
            "the sentence tokenizer a chance to split it into smaller pieces"
        )
        self.assertGreater(len(text), MAX_CHUNK_LENGTH)  # sanity check: fixture must trigger the split path

        result = chunk_text(text)

        self.assertEqual(result, [text])

    def test_normalizes_internal_whitespace(self):
        text = "This   paragraph has\nirregular   whitespace\nand line breaks  in the middle of it."

        result = chunk_text(text)

        self.assertEqual(
            result,
            ["This paragraph has irregular whitespace and line breaks in the middle of it."],
        )

    def test_trailing_short_sentence_merges_into_previous_group(self):
        first_sentence = (
            "This first sentence is carefully padded with extra filler words "
            "so it lands under the two hundred character limit today for sure."
        )
        second_sentence = (
            "This second standalone sentence is deliberately padded even further "
            "with additional filler words so that its own length alone comes in "
            "very close to the two hundred character chunk size limit used "
            "throughout this whole system right now."
        )
        trailing_sentence = "Yes it is."
        text = f"{first_sentence} {second_sentence} {trailing_sentence}"
        self.assertGreater(len(text), MAX_CHUNK_LENGTH)  # sanity check: fixture must trigger the split path

        result = chunk_text(text)

        self.assertEqual(len(result), 2)
        self.assertIn(first_sentence, result[0])
        self.assertIn(second_sentence, result[1])
        self.assertIn(trailing_sentence, result[1])
